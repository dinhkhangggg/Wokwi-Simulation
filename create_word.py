import docx
from docx.shared import Pt, Inches
import re

doc = docx.Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)

md_path = r'c:\Users\ADMIN\.gemini\antigravity\brain\c7956faa-8605-4145-915f-dac50e726124\chapter_4_evaluation.md'

with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

table_data = []
in_table = False

def process_table():
    global table_data, doc
    if not table_data: return
    rows = [r for r in table_data if '---' not in ''.join(r)]
    if rows:
        try:
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = 'Table Grid'
            for i, row_data in enumerate(rows):
                for j, cell_data in enumerate(row_data):
                    cell_text = cell_data.strip().replace('**', '').replace('$', '').replace('`', '')
                    table.cell(i, j).text = cell_text
        except Exception as e:
            print("Table error:", e)
    table_data = []

def add_paragraph_with_bold(p, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            p.add_run(part[2:-2]).bold = True
        else:
            part = part.replace('$$', '').replace('$', '').replace('_', ' ').replace('`', '')
            p.add_run(part)

for raw_line in lines:
    line = raw_line.strip()
    
    if line.startswith('|'):
        in_table = True
        # remove leading and trailing pipe
        content = line
        if content.startswith('|'): content = content[1:]
        if content.endswith('|'): content = content[:-1]
        cells = content.split('|')
        table_data.append(cells)
        continue
    else:
        if in_table:
            process_table()
            in_table = False
            
    if not line or line == '---':
        continue
        
    if line.startswith('#'):
        level = min(len(line.split(' ')[0]), 9)
        text = line.replace('#', '').strip()
        p = doc.add_heading(text, level=level)
    elif line.startswith('* '):
        text = line[2:].strip()
        try:
            p = doc.add_paragraph(style='List Bullet')
        except:
            p = doc.add_paragraph()
            p.add_run("• ")
        add_paragraph_with_bold(p, text)
    elif line.startswith('> '):
        text = line[2:].strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.add_run(text).italic = True
    else:
        p = doc.add_paragraph()
        add_paragraph_with_bold(p, line)

if in_table:
    process_table()

doc.save(r'c:\Users\ADMIN\Documents\AILINH\BÁO CÁO AIOT\prj\BaoCao_Chuong4.docx')
print("SUCCESS")
